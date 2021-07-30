# -*- coding: utf-8 -*-
# FileName : create_Doc.py 
# Author : xiaoran
# Date : 2021-07-28 15:22
import os

import docx

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式

# document = Document()
# p1 = document.add_paragraph()  # 初始化建立第一个自然段
# p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 对齐方式为居中，没有这句默认为左对齐
# run1 = p1.add_run('testxxxxxx')  # 添加文本
# run1.font.name = u'微软雅黑'  # 设置文本字体
# run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置字体支持中文
# run1.font.size = Pt(21)  # 设置字体大小为21磅
# run1.font.bold = True  # 设置加粗
# p1.space_after = Pt(5)  # 段后距离5磅
# p1.space_before = Pt(5)  # 段前距离5磅
# document.save('hello.docx')

list_tem = ['张三', '李四', '王武', '钱六', '狗蛋儿', '建国']
list_temp = ['张三', '李四']
save_path = "/Users/ranchuangyun/Documents/exercise/python/python_study/files_set/"
save_path1 = "/Users/ranchuangyun/Documents/exercise/python/python_study/tools/create_Doc.py"
if not os.path.exists(save_path):
    # os.makedirs(save_path)
    print('nonono')

for i in list_tem:
    docx_test = docx.Document()
    pp = docx_test.add_paragraph()
    font_set = pp.add_run(i)
    # font_set.font.name = u'微软雅黑'
    # font_set.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    # font_set.font.size = Pt(21)  # 设置字体大小为21磅
    # font_set.font.bold = True  # 设置加粗
    # pp.space_after = Pt(5)  # 段后距离5磅
    # pp.space_before = Pt(5)  # 段前距离5磅
    # print(i)

    docx_test.save("{}.docx".format(i))
